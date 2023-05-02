from flask import Flask, request, send_file, send_from_directory, jsonify
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
import xml.etree.ElementTree
import os
import shutil
import zipfile
import random
app = Flask(__name__)


def change(amount):
    # calculate the resultant change and store the result (res)
    res = []
    coins = [1, 5, 10, 25]  # value of pennies, nickels, dimes, quarters
    coin_lookup = {25: "quarters", 10: "dimes", 5: "nickels", 1: "pennies"}

    # divide the amount*100 (the amount in cents) by a coin value
    # record the number of coins that evenly divide and the remainder
    coin = coins.pop()
    num, rem = divmod(int(amount*100), coin)
    # append the coin type and number of coins that had no remainder
    res.append({num: coin_lookup[coin]})

    # while there is still some remainder, continue adding coins to the result
    while rem > 0:
        coin = coins.pop()
        num, rem = divmod(rem, coin)
        if num:
            if coin in coin_lookup:
                res.append({num: coin_lookup[coin]})
    return res


@app.route('/')
def hello():
    """Return a friendly HTTP greeting."""
    print("I am ihello world")
    return 'I make change and route /change'


@app.route('/change/<dollar>/<cents>')
def changeroute(dollar, cents):
    print(f"Make Change for {dollar}.{cents}")
    amount = f"{dollar}.{cents}"
    result = change(float(amount))
    return jsonify(result)


@app.route('/generateTest', methods=['POST'])
def testroute():
    # assigning request json to variables
    test = request.get_json()
    has_header = test['header']
    has_title = test['title']
    has_instructions = test['instructions']
    section = test['section']

    # document object
    question_doc = Document()
    answer_doc = Document()

    # generate header (queston & answer document)
    if has_header == True:
        header_left_alignment = test['header_left_alignment']
        header_text = test['entered_header']

        # question doc
        doc_settings = question_doc.sections[0]
        doc_settings.different_first_page_header_footer = True
        qd_header = doc_settings.first_page_header
        qd_header_para = qd_header.paragraphs[0]
        if header_left_alignment == False:
            qd_header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        qd_header_run = qd_header_para.add_run()
        qd_header_font = qd_header_run.font
        qd_header_font.size = Pt(12)
        qd_header_font.name = 'Helvetica'

        qd_header_run.text = header_text

        # answer doc
        doc_settings = answer_doc.sections[0]
        doc_settings.different_first_page_header_footer = True
        ad_header = doc_settings.first_page_header
        ad_header_para = ad_header.paragraphs[0]
        if header_left_alignment == False:
            ad_header_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ad_header_run = ad_header_para.add_run()
        ad_header_font = ad_header_run.font
        ad_header_font.size = Pt(12)
        ad_header_font.name = 'Helvetica'

        ad_header_run.text = header_text

    # generate title (question and answer document)
    if has_title == True:
        title_text = test['entered_title']

        # question doc
        qd_title_para = question_doc.add_paragraph()
        qd_title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qd_title_para.paragraph_format.space_before = Pt(12)
        qd_title_para.paragraph_format.space_after = Pt(12)
        qd_title_run = qd_title_para.add_run()
        qd_title_font = qd_title_run.font
        qd_title_font.name = 'Helvetica'
        qd_title_font.size = Pt(20)
        qd_title_font.bold = True

        qd_title_run.text = title_text

        # answer doc
        ad_title_para = answer_doc.add_paragraph()
        ad_title_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ad_title_para.paragraph_format.space_before = Pt(12)
        ad_title_para.paragraph_format.space_after = Pt(12)
        ad_title_run = ad_title_para.add_run()
        ad_title_font = ad_title_run.font
        ad_title_font.name = 'Helvetica'
        ad_title_font.size = Pt(20)
        ad_title_font.bold = True

        ad_title_run.text = title_text

    # generate instructions (question and answer document)
    if has_instructions == True:
        instructions_text = test['entered_instructions']

        # question doc
        qd_instructions_para = question_doc.add_paragraph()
        qd_instructions_para.paragraph_format.space_after = Pt(40)
        qd_instructions_run = qd_instructions_para.add_run()
        qd_instructions_font = qd_instructions_run.font
        qd_instructions_font.name = 'Helvetica'
        qd_instructions_font.size = Pt(12)

        qd_instructions_run.text = instructions_text

        # answer doc
        ad_instructions_para = answer_doc.add_paragraph()
        ad_instructions_para.paragraph_format.space_after = Pt(40)
        ad_instructions_run = ad_instructions_para.add_run()
        ad_instructions_font = ad_instructions_run.font
        ad_instructions_font.name = 'Helvetica'
        ad_instructions_font.size = Pt(12)

        ad_instructions_run.text = instructions_text

    # geneate section loop (question and answer document)
    TMP_QUESTION_INDEX = 10
    for i in range(len(section)):
        # generate section title
        has_section_title = section[i]['section_title']
        has_section_instructions = section[i]['section_instructions']
        question_type = section[i]['question_type']
        question_structure = section[i]['question_structure']

        if has_section_title:
            # question doc
            qd_section_title_para = question_doc.add_paragraph()
            qd_section_title_para.paragraph_format.space_before = Pt(20)
            qd_section_title_para.paragraph_format.space_after = Pt(0)
            qd_section_title_run = qd_section_title_para.add_run()
            qd_section_title_font = qd_section_title_run.font
            qd_section_title_font.name = 'Helvetica'
            qd_section_title_font.size = Pt(12)
            qd_section_title_font.bold = True

            qd_section_title_run.text = "section " + \
                f"{i + 1}: Multiple Choice"

            # answer doc
            ad_section_title_para = answer_doc.add_paragraph()
            ad_section_title_para.paragraph_format.space_before = Pt(20)
            ad_section_title_para.paragraph_format.space_after = Pt(0)
            ad_section_title_run = ad_section_title_para.add_run()
            ad_section_title_font = ad_section_title_run.font
            ad_section_title_font.name = 'Helvetica'
            ad_section_title_font.size = Pt(12)
            ad_section_title_font.bold = True

            ad_section_title_run.text = "section " + \
                f"{i + 1}: Multiple Choice"

        # generate section instructions (question and answer doc)
        if has_section_instructions:
            # question doc
            qd_section_instructions_para = question_doc.add_paragraph()
            qd_section_instructions_para.paragraph_format.space_before = Pt(0)
            qd_section_instructions_para.paragraph_format.space_after = Pt(0)
            qd_section_instructions_run = qd_section_instructions_para.add_run()
            qd_section_instructions_font = qd_section_instructions_run.font
            qd_section_instructions_font.name = 'Helvetica'
            qd_section_instructions_font.size = Pt(12)

            qd_section_instructions_run.text = "Circle the correct answer"

            # answer doc
            ad_section_instructions_para = question_doc.add_paragraph()
            ad_section_instructions_para.paragraph_format.space_before = Pt(0)
            ad_section_instructions_para.paragraph_format.space_after = Pt(0)
            ad_section_instructions_run = ad_section_instructions_para.add_run()
            ad_section_instructions_font = ad_section_instructions_run.font
            ad_section_instructions_font.name = 'Helvetica'
            ad_section_instructions_font.size = Pt(12)

            ad_section_instructions_run.text = "Circle the correct answer"

        # generate questions (question and answer doc)
        if question_type == "MC":
            questions = question_structure['questions']
            for j in range(len(questions)):
                question_text = questions[j]['entered_question']

                # question doc
                qd_para = question_doc.add_paragraph()

                qd_para.paragraph_format.space_before = Pt(12)
                qd_para.paragraph_format.space_after = Pt(12)
                qd_run = qd_para.add_run()
                qd_font = qd_run.font
                qd_font.name = 'Helvetica'
                qd_font.size = Pt(12)

                qd_run.text = f"{j + 1}. {question_text}\n"

                # answer doc
                ad_para = answer_doc.add_paragraph()

                ad_para.paragraph_format.space_before = Pt(12)
                ad_para.paragraph_format.space_after = Pt(12)
                ad_run = ad_para.add_run()
                ad_font = ad_run.font
                ad_font.name = 'Helvetica'
                ad_font.size = Pt(12)

                ad_run.text = f"{j + 1}. {question_text}\n"

                answer_options = questions[j]['answer_options']
                correct_answer_id = questions[j]['correct_answer_id']
                for k in range(len(answer_options)):
                    answer_option_text = answer_options[k]['entered_option']
                    option_letter = chr(97 + k)
                    # question document
                    qd_run.text += f"    {option_letter}) {answer_option_text}\n"

                    # answer document
                    ad_run.text += f"    {option_letter}) {answer_option_text}\n"

                    if(answer_options[k]['id'] == correct_answer_id):
                        border_xml = '<w:bdr w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                        border_el = parse_xml(border_xml)[0]
                        ad_run.append(border_el)

    mainPath = Path.home() / 'Exam'
    if(not(os.path.isdir(mainPath))):
        os.mkdir(mainPath)

    testPath = mainPath / 'TestPath'
    if(not(os.path.isdir(testPath))):
        os.mkdir(testPath)

    answerPath = mainPath/'AnswerPath'
    if(not(os.path.isdir(answerPath))):
        os.mkdir(answerPath)

    os.chdir(testPath)
    question_doc.save('test.docx')

    os.chdir(answerPath)
    answer_doc.save('answer_key.docx')

    os.chdir(Path.home())

    zipFile = shutil.make_archive('Exam', 'zip', Path.home() / 'Exam')

    return (
        send_file(
            zipFile,
            mimetype='zip',
            download_name='Exam.zip',
            as_attachment=True
        )
    )


if __name__ == '__main__':
    port = os.environ.get("PORT", 4000)
    app.run(debug=True, host='0.0.0.0', port=port)
